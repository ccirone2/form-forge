"""Tests for FormForge DOCX templates."""

import json
import sys
import importlib.util
from io import BytesIO
from pathlib import Path

from docx import Document

FIXTURES = Path(__file__).resolve().parent / "fixtures"
TEMPLATES = Path(__file__).resolve().parent.parent / "templates"

# Add templates/ to path so `import stencils` works
sys.path.insert(0, str(TEMPLATES))


def load_template(name):
    spec = importlib.util.spec_from_file_location(name, TEMPLATES / f"{name}.py")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def _full_text(docx_bytes):
    """Read DOCX bytes back and extract all text from paragraphs and table cells."""
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
#  Smoke tests (valid ZIP / non-empty)
# ---------------------------------------------------------------------------


def test_onboarding_generates_valid_docx():
    data = json.loads((FIXTURES / "onboarding_sample.json").read_text())
    mod = load_template("onboarding")
    result = mod.generate_docx(data)
    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result[:2] == b"PK"


def test_onboarding_with_empty_data():
    mod = load_template("onboarding")
    result = mod.generate_docx({})
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"


def test_expense_report_generates_valid_docx():
    data = json.loads((FIXTURES / "expense-report_sample.json").read_text())
    mod = load_template("expense-report")
    result = mod.generate_docx(data)
    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result[:2] == b"PK"


def test_expense_report_with_empty_data():
    mod = load_template("expense-report")
    result = mod.generate_docx({})
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"


def test_field_type_demo_generates_valid_docx():
    data = json.loads((FIXTURES / "field-type-demo_sample.json").read_text())
    mod = load_template("field-type-demo")
    result = mod.generate_docx(data)
    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result[:2] == b"PK"


def test_field_type_demo_with_empty_data():
    mod = load_template("field-type-demo")
    result = mod.generate_docx({})
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"


# ---------------------------------------------------------------------------
#  Round-trip content assertions
# ---------------------------------------------------------------------------


def test_onboarding_content_round_trip():
    data = json.loads((FIXTURES / "onboarding_sample.json").read_text())
    mod = load_template("onboarding")
    result = mod.generate_docx(data)
    text = _full_text(result)
    assert "Alice" in text
    assert "Johnson" in text
    assert "Engineering" in text
    assert "Senior Developer" in text


def test_expense_report_content_round_trip():
    data = json.loads((FIXTURES / "expense-report_sample.json").read_text())
    mod = load_template("expense-report")
    result = mod.generate_docx(data)
    text = _full_text(result)
    assert "Jane Doe" in text
    assert "Engineering" in text
    assert "Flight to NYC" in text
    assert "$1,375.50" in text


def test_field_type_demo_content_round_trip():
    data = json.loads((FIXTURES / "field-type-demo_sample.json").read_text())
    mod = load_template("field-type-demo")
    result = mod.generate_docx(data)
    text = _full_text(result)
    assert "Jane Doe" in text
    assert "Conference" in text
    assert "In-person" in text
    assert "100 Convention Center Dr" in text


# ---------------------------------------------------------------------------
#  visible_when hidden-field path
# ---------------------------------------------------------------------------


def test_field_type_demo_virtual_path():
    """When attendance_mode is Virtual, venue/dietary are hidden (empty)
    and virtual_platform appears in the output."""
    data = json.loads((FIXTURES / "field-type-demo_virtual_sample.json").read_text())
    mod = load_template("field-type-demo")
    result = mod.generate_docx(data)
    text = _full_text(result)
    assert "Zoom" in text
    assert "100 Convention Center Dr" not in text
    assert "Vegetarian" not in text
