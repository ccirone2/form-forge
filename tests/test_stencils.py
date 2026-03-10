"""Tests for templates/stencils.py shared utilities."""

import sys
import zipfile
from pathlib import Path

import pytest
from docx.shared import RGBColor, Inches
from docx.oxml.ns import qn
from lxml import etree

# Add templates/ to path so `import stencils` works
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "templates"))

import stencils


# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx_styles_xml(doc):
    """Save doc to bytes and return parsed styles.xml root."""
    import io

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with zipfile.ZipFile(buf) as z:
        return etree.fromstring(z.read("word/styles.xml"))


def _table_border_vals(table):
    """Return a dict of border edge -> val from the table's XML."""
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None
    if borders is None:
        return {}
    result = {}
    for child in borders:
        tag = child.tag.split("}")[-1]
        result[tag] = child.get(qn("w:val"))
    return result


# ---------------------------------------------------------------------------
#  new_doc
# ---------------------------------------------------------------------------


def test_new_doc_creates_document():
    doc = stencils.new_doc("Test Title")
    assert doc is not None
    assert len(doc.paragraphs) > 0


def test_new_doc_with_subtitle():
    doc = stencils.new_doc("Test Title", "Test Subtitle")
    texts = [p.text for p in doc.paragraphs]
    assert "Test Subtitle" in texts


def test_new_doc_default_font():
    doc = stencils.new_doc("Title")
    style = doc.styles["Normal"]
    assert style.font.name == "Segoe UI"
    assert style.font.size.pt == 11


def test_new_doc_custom_font():
    doc = stencils.new_doc("Title", font_name="Arial", font_size=12)
    style = doc.styles["Normal"]
    assert style.font.name == "Arial"
    assert style.font.size.pt == 12


# ---------------------------------------------------------------------------
#  Theme tests (replacing palette tests)
# ---------------------------------------------------------------------------


def test_theme_classic_has_all_fields():
    t = stencils.THEME_CLASSIC
    for field in stencils._THEME_FIELDS:
        assert hasattr(t, field), f"Missing field: {field}"


def test_theme_minimal_has_all_fields():
    t = stencils.THEME_MINIMAL
    for field in stencils._THEME_FIELDS:
        assert hasattr(t, field), f"Missing field: {field}"


def test_theme_modern_has_all_fields():
    t = stencils.THEME_MODERN
    for field in stencils._THEME_FIELDS:
        assert hasattr(t, field), f"Missing field: {field}"


def test_classic_theme_matches_legacy_colors():
    """THEME_CLASSIC colors must match the original PALETTE_CLASSIC values."""
    assert stencils.THEME_CLASSIC.color_title == RGBColor(0x33, 0x33, 0x66)
    assert stencils.THEME_CLASSIC.color_subtitle == RGBColor(0x66, 0x66, 0x99)
    assert stencils.THEME_CLASSIC.color_muted == RGBColor(0x99, 0x99, 0x99)
    assert stencils.THEME_CLASSIC.color_footer == RGBColor(0xAA, 0xAA, 0xAA)
    assert stencils.THEME_CLASSIC.color_accent == RGBColor(0x1A, 0x1A, 0x3E)


def test_set_theme_switches_title_style_color():
    try:
        stencils.set_theme(stencils.THEME_MINIMAL)
        doc = stencils.new_doc("Test")
        assert doc.styles["Title"].font.color.rgb == stencils.THEME_MINIMAL.color_title
    finally:
        stencils.set_theme(stencils.THEME_MODERN)


def test_set_theme_invalid_raises():
    with pytest.raises(ValueError, match="missing required fields"):
        stencils.set_theme(object())


def test_set_theme_custom_theme():
    custom = stencils.DocTheme(
        color_title=RGBColor(0xFF, 0x00, 0x00),
        color_subtitle=RGBColor(0x00, 0xFF, 0x00),
        color_muted=RGBColor(0x00, 0x00, 0xFF),
        color_footer=RGBColor(0xAA, 0xBB, 0xCC),
        color_accent=RGBColor(0x11, 0x22, 0x33),
        font_body="Arial",
        font_heading="Arial Bold",
        font_caption="Arial Narrow",
        size_body=12,
        size_title=28,
        size_heading1=18,
        size_heading2=14,
        size_heading3=12,
        size_heading4=11,
        size_heading5=11,
        size_heading6=11,
        size_subtitle=13,
        size_table=11,
        size_caption=10,
        size_footer=9,
        margin_top=0.5,
        margin_bottom=0.5,
        margin_left=0.75,
        margin_right=0.75,
    )
    try:
        stencils.set_theme(custom)
        doc = stencils.new_doc("Test")
        assert doc.styles["Title"].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
        assert doc.styles["Normal"].font.name == "Arial"
    finally:
        stencils.set_theme(stencils.THEME_MODERN)


def test_set_theme_restores_default():
    stencils.set_theme(stencils.THEME_MINIMAL)
    stencils.set_theme(stencils.THEME_MODERN)
    assert stencils._active_theme is stencils.THEME_MODERN


def test_new_doc_theme_override():
    """theme= on new_doc overrides styles without changing active theme."""
    before = stencils._active_theme
    doc = stencils.new_doc("Test", theme=stencils.THEME_MINIMAL)
    assert doc.styles["Title"].font.color.rgb == stencils.THEME_MINIMAL.color_title
    assert stencils._active_theme is before


def test_new_doc_theme_override_does_not_change_active():
    before = stencils._active_theme
    stencils.new_doc("Test", theme=stencils.THEME_CLASSIC)
    assert stencils._active_theme is before


def test_new_doc_subtitle_theme_override():
    doc = stencils.new_doc("T", "Sub", theme=stencils.THEME_MODERN)
    assert doc.styles["Subtitle"].font.color.rgb == stencils.THEME_MODERN.color_subtitle


def test_new_doc_no_theme_uses_active():
    try:
        stencils.set_theme(stencils.THEME_MINIMAL)
        doc = stencils.new_doc("Test")
        assert doc.styles["Title"].font.color.rgb == stencils.THEME_MINIMAL.color_title
    finally:
        stencils.set_theme(stencils.THEME_MODERN)


# ---------------------------------------------------------------------------
#  Template builder
# ---------------------------------------------------------------------------


def test_template_bytes_are_valid_docx():
    raw = stencils._get_template(stencils.THEME_MODERN)
    assert isinstance(raw, bytes)
    assert raw[:2] == b"PK"


def test_template_has_only_normal_table_style():
    doc = stencils.new_doc("Test")
    sroot = _docx_styles_xml(doc)
    table_styles = [
        s
        for s in sroot.findall(f"{{{WNS}}}style")
        if s.get(f"{{{WNS}}}type") == "table"
    ]
    assert len(table_styles) == 1
    name_el = table_styles[0].find(f"{{{WNS}}}name")
    assert name_el.get(f"{{{WNS}}}val") == "Normal Table"


def test_template_heading1_style_configured():
    doc = stencils.new_doc("Test", theme=stencils.THEME_CLASSIC)
    h1 = doc.styles["Heading 1"]
    assert h1.font.name == "Segoe UI Semibold"
    assert h1.font.color.rgb == stencils.THEME_CLASSIC.color_title


def test_template_heading2_style_configured():
    doc = stencils.new_doc("Test", theme=stencils.THEME_CLASSIC)
    h2 = doc.styles["Heading 2"]
    assert h2.font.name == "Segoe UI Semibold"
    assert h2.font.color.rgb == stencils.THEME_CLASSIC.color_subtitle


def test_template_page_margins():
    doc = stencils.new_doc("Test")
    section = doc.sections[0]
    assert section.top_margin == Inches(1.0)
    assert section.bottom_margin == Inches(0.75)
    assert section.left_margin == Inches(1.5)
    assert section.right_margin == Inches(1.5)


def test_template_custom_margins():
    custom = stencils.DocTheme(
        color_title=RGBColor(0, 0, 0),
        color_subtitle=RGBColor(0, 0, 0),
        color_muted=RGBColor(0, 0, 0),
        color_footer=RGBColor(0, 0, 0),
        color_accent=RGBColor(0, 0, 0),
        font_body="Arial",
        font_heading="Arial",
        font_caption="Arial",
        size_body=11,
        size_title=26,
        size_heading1=16,
        size_heading2=13,
        size_heading3=12,
        size_heading4=11,
        size_heading5=11,
        size_heading6=11,
        size_subtitle=12,
        size_table=10,
        size_caption=9,
        size_footer=8,
        margin_top=0.5,
        margin_bottom=0.75,
        margin_left=1.25,
        margin_right=1.5,
    )
    doc = stencils.new_doc("Test", theme=custom)
    section = doc.sections[0]
    assert section.top_margin == Inches(0.5)
    assert section.bottom_margin == Inches(0.75)
    assert section.left_margin == Inches(1.25)
    assert section.right_margin == Inches(1.5)


def test_set_theme_rebuilds_cache():
    """Changing theme should produce template with new colors."""
    try:
        stencils.set_theme(stencils.THEME_CLASSIC)
        doc = stencils.new_doc("Test")
        assert doc.styles["Title"].font.color.rgb == stencils.THEME_CLASSIC.color_title

        stencils.set_theme(stencils.THEME_MINIMAL)
        doc = stencils.new_doc("Test")
        assert doc.styles["Title"].font.color.rgb == stencils.THEME_MINIMAL.color_title
    finally:
        stencils.set_theme(stencils.THEME_MODERN)


# ---------------------------------------------------------------------------
#  Table styling
# ---------------------------------------------------------------------------


def test_table_section():
    doc = stencils.new_doc("Test")
    stencils.table_section(
        doc,
        "Info",
        [("Name", "Alice"), ("Role", "Developer")],
    )
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.rows[0].cells[0].text == "Name"
    assert table.rows[0].cells[1].text == "Alice"


def test_table_section_empty_value():
    doc = stencils.new_doc("Test")
    stencils.table_section(doc, "Info", [("Field", "")])
    table = doc.tables[0]
    assert table.rows[0].cells[1].text == "\u2014"


def test_table_section_borders_are_none():
    doc = stencils.new_doc("Test")
    stencils.table_section(doc, "Info", [("A", "B")])
    borders = _table_border_vals(doc.tables[0])
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        assert borders.get(edge) == "none", f"{edge} should be 'none'"


def test_table_section_zero_value():
    doc = stencils.new_doc("Test")
    stencils.table_section(doc, "Info", [("Count", "0")])
    table = doc.tables[0]
    assert table.rows[0].cells[1].text == "0"


def test_table_section_numeric_zero():
    doc = stencils.new_doc("Test")
    stencils.table_section(doc, "Info", [("Count", 0)])
    table = doc.tables[0]
    assert table.rows[0].cells[1].text == "0"


def test_repeater_table_with_items():
    doc = stencils.new_doc("Test")
    items = [
        {"desc": "Item A", "amount": "100.50"},
        {"desc": "Item B", "amount": "200"},
    ]
    stencils.repeater_table(
        doc,
        headers=["Description", "Amount"],
        items=items,
        field_keys=["desc", "amount"],
        currency_keys=["amount"],
    )
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert table.rows[1].cells[0].text == "Item A"
    assert table.rows[1].cells[1].text == "$100.50"
    assert table.rows[2].cells[1].text == "$200.00"


def test_repeater_table_empty():
    doc = stencils.new_doc("Test")
    stencils.repeater_table(
        doc,
        headers=["A", "B"],
        items=[],
        field_keys=["a", "b"],
    )
    texts = [p.text for p in doc.paragraphs]
    assert any("No line items provided." in t for t in texts)


def test_repeater_table_has_grid_borders():
    doc = stencils.new_doc("Test")
    stencils.repeater_table(
        doc,
        headers=["A"],
        items=[{"a": "1"}],
        field_keys=["a"],
    )
    borders = _table_border_vals(doc.tables[0])
    for edge in ("top", "bottom", "insideH"):
        assert borders.get(edge) == "single", f"{edge} should be 'single'"
    for edge in ("left", "right", "insideV"):
        assert borders.get(edge) == "none", f"{edge} should be 'none'"


def test_repeater_table_header_is_shaded():
    doc = stencils.new_doc("Test")
    stencils.repeater_table(
        doc,
        headers=["A"],
        items=[{"a": "1"}],
        field_keys=["a"],
    )
    header_cell = doc.tables[0].rows[0].cells[0]
    tcPr = header_cell._tc.tcPr
    shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
    assert shd is not None
    assert shd.get(qn("w:fill")) == str(stencils.THEME_MODERN.color_accent)


# ---------------------------------------------------------------------------
#  Other stencil functions
# ---------------------------------------------------------------------------


def test_longtext_with_content():
    doc = stencils.new_doc("Test")
    stencils.longtext(doc, "Notes", "Line one\nLine two")
    texts = [p.text for p in doc.paragraphs]
    assert "Line one" in texts
    assert "Line two" in texts


def test_longtext_empty():
    doc = stencils.new_doc("Test")
    stencils.longtext(doc, "Notes", "")
    texts = [p.text for p in doc.paragraphs]
    assert "No information provided." in texts


def test_bullet_list_with_items():
    doc = stencils.new_doc("Test")
    stencils.bullet_list(doc, "Skills", "Python\nJavaScript\nRust")
    texts = [p.text for p in doc.paragraphs]
    assert "Python" in texts
    assert "JavaScript" in texts
    assert "Rust" in texts


def test_bullet_list_empty():
    doc = stencils.new_doc("Test")
    stencils.bullet_list(doc, "Skills", "")
    texts = [p.text for p in doc.paragraphs]
    assert "No items listed." in texts


def test_signatures():
    doc = stencils.new_doc("Test")
    stencils.signatures(doc, ["Employee", "Date", "Manager", "Date"])
    assert len(doc.tables) >= 1
    sig_table = doc.tables[-1]
    assert len(sig_table.rows) == 2
    assert len(sig_table.columns) == 2


def test_signatures_odd_count():
    doc = stencils.new_doc("Test")
    stencils.signatures(doc, ["Solo Signer"])
    sig_table = doc.tables[-1]
    assert len(sig_table.rows) == 1


def test_finalize_returns_valid_docx():
    doc = stencils.new_doc("Test")
    result = stencils.finalize(doc)
    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result[:2] == b"PK"


def test_footer():
    doc = stencils.new_doc("Test")
    stencils.footer(doc)
    texts = [p.text for p in doc.paragraphs]
    assert any("auto-generated by FormForge" in t for t in texts)


def test_address_full():
    doc = stencils.new_doc("Test")
    import json

    addr = json.dumps(
        {"street": "123 Main St", "city": "Springfield", "state": "IL", "zip": "62701"}
    )
    stencils.address(doc, "Address", addr)
    texts = [p.text for p in doc.paragraphs]
    assert any("123 Main St" in t for t in texts)
    assert any("Springfield, IL 62701" in t for t in texts)


def test_address_empty_city():
    doc = stencils.new_doc("Test")
    import json

    addr = json.dumps(
        {"street": "123 Main St", "city": "", "state": "IL", "zip": "62701"}
    )
    stencils.address(doc, "Address", addr)
    texts = [p.text for p in doc.paragraphs]
    for t in texts:
        assert not t.strip().startswith(",")


def test_address_no_street():
    doc = stencils.new_doc("Test")
    import json

    addr = json.dumps(
        {"street": "", "city": "Springfield", "state": "IL", "zip": "62701"}
    )
    stencils.address(doc, "Address", addr)
    texts = [p.text for p in doc.paragraphs]
    assert any("No address provided." in t for t in texts)


def test_address_empty_json():
    doc = stencils.new_doc("Test")
    stencils.address(doc, "Address", "{}")
    texts = [p.text for p in doc.paragraphs]
    assert any("No address provided." in t for t in texts)


def test_image_empty():
    doc = stencils.new_doc("Test")
    stencils.image(doc, "", placeholder="No image.")
    texts = [p.text for p in doc.paragraphs]
    assert any("No image." in t for t in texts)


def test_image_invalid_b64():
    doc = stencils.new_doc("Test")
    stencils.image(doc, "data:image/png;base64,NOT_VALID!!!", placeholder="Failed.")
    texts = [p.text for p in doc.paragraphs]
    assert any("Failed." in t for t in texts)


def test_signature_empty():
    doc = stencils.new_doc("Test")
    stencils.signature(doc, "", "Employee Signature")
    texts = [p.text for p in doc.paragraphs]
    assert any("Employee Signature" in t for t in texts)
    assert any("_" * 40 in t for t in texts)


def test_signature_invalid_b64():
    doc = stencils.new_doc("Test")
    stencils.signature(doc, "data:image/png;base64,BAD_DATA!!!", "Signer")
    texts = [p.text for p in doc.paragraphs]
    assert any("Signer" in t for t in texts)
    assert any("_" * 40 in t for t in texts)
