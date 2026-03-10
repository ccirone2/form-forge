"""
FormForge Shared Template Stencils
===================================
Shared helper functions for all FormForge DOCX templates.
Works in both standard Python (import stencils) and Pyodide (via exec()).

Usage in templates:
    import stencils
    doc = stencils.new_doc("My Form Title", "Subtitle here")
    stencils.table_section(doc, "Section Name", [("Label", "Value"), ...])
    stencils.longtext(doc, "Notes", some_text)
    stencils.bullet_list(doc, "Skills", newline_separated_str)
    stencils.signatures(doc, ["Signer 1", "Signer 2"])
    stencils.footer(doc)
    return stencils.finalize(doc)
"""

import io
import json
import base64
import binascii
from dataclasses import dataclass
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
#  Theme system
# ---------------------------------------------------------------------------

_THEME_FIELDS = (
    "color_title",
    "color_subtitle",
    "color_muted",
    "color_footer",
    "color_accent",
    "font_body",
    "font_heading",
    "font_caption",
    "size_body",
    "size_title",
    "size_heading1",
    "size_heading2",
    "size_heading3",
    "size_heading4",
    "size_heading5",
    "size_heading6",
    "size_subtitle",
    "size_table",
    "size_caption",
    "size_footer",
    "margin_top",
    "margin_bottom",
    "margin_left",
    "margin_right",
)


@dataclass(frozen=True)
class DocTheme:
    """Complete document theme: colors, fonts, sizes, and page layout.

    Colors:
        color_title:    Main heading / Heading 1 text color.
        color_subtitle: Subtitle / Heading 2 text color.
        color_muted:    Empty-state / placeholder text color.
        color_footer:   Footer text color.
        color_accent:   Table header background / heavy emphasis.

    Fonts:
        font_body:    Base body font.
        font_heading: Heading / label font (semibold weight).
        font_caption: Small print / caption font (light weight).

    Sizes (points):
        size_body:     Normal body text.
        size_title:    Title heading (level 0).
        size_heading1: Level-1 heading.
        size_heading2: Level-2 heading.
        size_heading3: Level-3 heading.
        size_heading4: Level-4 heading.
        size_heading5: Level-5 heading.
        size_heading6: Level-6 heading.
        size_subtitle: Subtitle text.
        size_table:    Table cell / section body text.
        size_caption:  Signature labels / small text.
        size_footer:   Footer text.

    Page layout (inches):
        margin_top, margin_bottom, margin_left, margin_right.
    """

    # Colors
    color_title: RGBColor
    color_subtitle: RGBColor
    color_muted: RGBColor
    color_footer: RGBColor
    color_accent: RGBColor
    # Fonts
    font_body: str
    font_heading: str
    font_caption: str
    # Sizes (pt)
    size_body: int
    size_title: int
    size_heading1: int
    size_heading2: int
    size_heading3: int
    size_heading4: int
    size_heading5: int
    size_heading6: int
    size_subtitle: int
    size_table: int
    size_caption: int
    size_footer: int
    # Page layout (inches)
    margin_top: float
    margin_bottom: float
    margin_left: float
    margin_right: float


# Built-in themes ──────────────────────────────────────────────────────────

THEME_CLASSIC = DocTheme(
    color_title=RGBColor(0x33, 0x33, 0x66),
    color_subtitle=RGBColor(0x66, 0x66, 0x99),
    color_muted=RGBColor(0x99, 0x99, 0x99),
    color_footer=RGBColor(0xAA, 0xAA, 0xAA),
    color_accent=RGBColor(0x1A, 0x1A, 0x3E),
    font_body="Segoe UI",
    font_heading="Segoe UI Semibold",
    font_caption="Segoe UI Semilight",
    size_body=11,
    size_title=26,
    size_heading1=16,
    size_heading2=13,
    size_heading3=12,
    size_heading4=11,
    size_heading5=11,
    size_heading6=11,
    size_subtitle=12,
    size_table=10,
    size_caption=9,
    size_footer=8,
    margin_top=1.0,
    margin_bottom=1.0,
    margin_left=1.0,
    margin_right=1.0,
)

THEME_MINIMAL = DocTheme(
    color_title=RGBColor(0x1A, 0x1A, 0x1A),
    color_subtitle=RGBColor(0x55, 0x55, 0x55),
    color_muted=RGBColor(0xAA, 0xAA, 0xAA),
    color_footer=RGBColor(0xCC, 0xCC, 0xCC),
    color_accent=RGBColor(0x00, 0x00, 0x00),
    font_body="Segoe UI",
    font_heading="Segoe UI Semibold",
    font_caption="Segoe UI Semilight",
    size_body=11,
    size_title=26,
    size_heading1=16,
    size_heading2=13,
    size_heading3=12,
    size_heading4=11,
    size_heading5=11,
    size_heading6=11,
    size_subtitle=12,
    size_table=10,
    size_caption=9,
    size_footer=8,
    margin_top=1.0,
    margin_bottom=1.0,
    margin_left=1.0,
    margin_right=1.0,
)

THEME_MODERN = DocTheme(
    color_title=RGBColor(0x1B, 0x5E, 0x6E),
    color_subtitle=RGBColor(0x4A, 0x8F, 0xA3),
    color_muted=RGBColor(0x8F, 0xA9, 0xB2),
    color_footer=RGBColor(0xB0, 0xC4, 0xCB),
    color_accent=RGBColor(0x0D, 0x3D, 0x4A),
    font_body="Segoe UI",
    font_heading="Segoe UI Semibold",
    font_caption="Segoe UI Semilight",
    size_body=11,
    size_title=26,
    size_heading1=16,
    size_heading2=13,
    size_heading3=12,
    size_heading4=11,
    size_heading5=11,
    size_heading6=11,
    size_subtitle=12,
    size_table=10,
    size_caption=9,
    size_footer=8,
    margin_top=1.0,
    margin_bottom=0.75,
    margin_left=1.5,
    margin_right=1.5,
)

_active_theme = THEME_MODERN


def set_theme(theme):
    """Set the active document theme for all subsequent stencils calls.

    Args:
        theme: A DocTheme instance (use a built-in THEME_* constant or
               construct a custom DocTheme with all required fields).

    Raises:
        ValueError: If theme is missing any required field.
    """
    global _active_theme

    missing = [f for f in _THEME_FIELDS if not hasattr(theme, f)]
    if missing:
        raise ValueError(f"Theme missing required fields: {', '.join(missing)}")

    _active_theme = theme


# ---------------------------------------------------------------------------
#  Template builder
# ---------------------------------------------------------------------------

_template_cache = {}


def _set_style_font(style, font_name):
    """Set font on a style and strip theme-linked font attributes.

    Word's built-in styles (Title, Heading 1-4, Subtitle, etc.) carry
    ``w:asciiTheme`` / ``w:hAnsiTheme`` attributes on their ``rFonts``
    element.  These theme references take priority over explicit font
    names, causing python-docx's ``style.font.name = ...`` to be ignored
    when the document is opened in Word.  This helper removes the theme
    attributes so the explicit font actually takes effect.
    """
    style.font.name = font_name
    rPr = style.element.rPr
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
                rFonts.attrib.pop(qn(attr), None)


def _build_template(theme):
    """Build a DOCX template with all Word styles configured from a theme.

    The template has no body content — just style definitions, page layout,
    and a clean style sheet with no built-in table styles.  ``new_doc()``
    clones from the cached bytes so every document inherits the correct
    headings, fonts, and margins automatically.
    """
    doc = Document()

    # -- Strip all table styles except Normal Table --------------------
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    styles_el = doc.styles.element
    for style_el in styles_el.findall(f"{{{ns}}}style"):
        if style_el.get(f"{{{ns}}}type") == "table" and not style_el.get(f"{{{ns}}}default"):
            styles_el.remove(style_el)

    # -- Configure paragraph styles ------------------------------------
    normal = doc.styles["Normal"]
    _set_style_font(normal, theme.font_body)
    normal.font.size = Pt(theme.size_body)
    normal.paragraph_format.space_after = Pt(0)

    title_style = doc.styles["Title"]
    _set_style_font(title_style, theme.font_heading)
    title_style.font.size = Pt(theme.size_title)
    title_style.font.color.rgb = theme.color_title
    title_style.font.bold = False
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, color in [
        ("Heading 1", theme.size_heading1, theme.color_title),
        ("Heading 2", theme.size_heading2, theme.color_subtitle),
    ]:
        h = doc.styles[name]
        _set_style_font(h, theme.font_heading)
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = False

    for name, size, color in [
        ("Heading 3", theme.size_heading3, theme.color_title),
        ("Heading 4", theme.size_heading4, theme.color_subtitle),
        ("Heading 5", theme.size_heading5, theme.color_subtitle),
        ("Heading 6", theme.size_heading6, theme.color_subtitle),
    ]:
        h = doc.styles[name]
        _set_style_font(h, theme.font_heading)
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = False

    try:
        sub = doc.styles["Subtitle"]
        _set_style_font(sub, theme.font_caption)
        sub.font.size = Pt(theme.size_subtitle)
        sub.font.color.rgb = theme.color_subtitle
        sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except KeyError:
        pass  # Handled per-run in new_doc() as fallback

    try:
        lb = doc.styles["List Bullet"]
        _set_style_font(lb, theme.font_body)
        lb.font.size = Pt(theme.size_table)
    except KeyError:
        pass

    # -- Page layout ---------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(theme.margin_top)
        section.bottom_margin = Inches(theme.margin_bottom)
        section.left_margin = Inches(theme.margin_left)
        section.right_margin = Inches(theme.margin_right)

    # -- Clear body content (keep sectPr) ------------------------------
    body = doc.element.body
    for p in body.findall(qn("w:p")):
        body.remove(p)

    # -- Serialize -----------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _get_template(theme):
    """Return cached template bytes for a theme, building if needed."""
    if theme not in _template_cache:
        _template_cache[theme] = _build_template(theme)
    return _template_cache[theme]


# Pre-build the default theme's template at import time
_template_cache[_active_theme] = _build_template(_active_theme)


# ---------------------------------------------------------------------------
#  Table formatting helpers
# ---------------------------------------------------------------------------


def _set_cell_margins(table, top=None, bottom=None):
    """Set default top/bottom cell margins for all cells in a table (inches)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    existing = tblPr.find(qn("w:tblCellMar"))
    if existing is not None:
        tblPr.remove(existing)
    parts = [f"<w:tblCellMar {nsdecls('w')}>"]
    if top is not None:
        twips = int(top * 1440)
        parts.append(f'<w:top w:w="{twips}" w:type="dxa"/>')
    if bottom is not None:
        twips = int(bottom * 1440)
        parts.append(f'<w:bottom w:w="{twips}" w:type="dxa"/>')
    parts.append("</w:tblCellMar>")
    tblPr.append(parse_xml("".join(parts)))


def _set_table_borders(table, val="single", sz=4, color="BFBFBF"):
    """Apply uniform borders to every edge of a table (or remove them)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = parse_xml(
        f"<w:tblBorders {nsdecls('w')}>"
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    tblPr.append(borders)


def _shade_cells(row, fill_hex):
    """Apply background shading to every cell in a table row."""
    for cell in row.cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
#  Public API
# ---------------------------------------------------------------------------


def new_doc(title_text, subtitle_text="", font_name=None, font_size=None, theme=None):
    """
    Create a styled Document with a centered title and optional subtitle.

    The document is cloned from a pre-built template whose heading styles,
    page margins, and base font are configured from the active (or given)
    theme.

    Args:
        title_text: Main heading text.
        subtitle_text: Optional subtitle displayed below the title.
        font_name: Override the base body font (default: theme's font_body).
        font_size: Override the base font size in points (default: theme's
                   size_body).
        theme: Optional DocTheme for this document only.
               If None, uses the current active theme.

    Returns:
        A python-docx Document instance.
    """
    t = theme if theme is not None else _active_theme
    doc = Document(io.BytesIO(_get_template(t)))

    if font_name is not None or font_size is not None:
        style = doc.styles["Normal"]
        if font_name is not None:
            style.font.name = font_name
        if font_size is not None:
            style.font.size = Pt(font_size)

    doc.add_heading(title_text, level=0)
    doc.add_paragraph("")

    if subtitle_text:
        try:
            doc.add_paragraph(subtitle_text, style="Subtitle")
        except KeyError:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(subtitle_text)
            run.font.name = t.font_caption
            run.font.size = Pt(t.size_subtitle)
            run.font.color.rgb = t.color_subtitle
        doc.add_paragraph("")

    return doc


def table_section(doc, heading, rows):
    """
    Add a heading followed by a borderless two-column key/value table.
    Labels in the first column use the heading font (semibold).

    Args:
        doc: The Document instance.
        heading: Section heading string.
        rows: List of (label, value) tuples.
    """
    t = _active_theme
    doc.add_heading(heading, level=1)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, val="none", sz=0, color="auto")
    _set_cell_margins(table, top=0.02, bottom=0.02)

    for label, value in rows:
        row = table.add_row()

        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.font.name = t.font_heading
        lr.font.size = Pt(t.size_table)

        vp = row.cells[1].paragraphs[0]
        display_value = str(value) if value is not None and value != "" else "\u2014"
        vr = vp.add_run(display_value)
        vr.font.size = Pt(t.size_table)

    doc.add_paragraph("")


def longtext(doc, heading, text):
    """
    Add a sub-heading followed by one or more paragraphs of body text.
    Respects newline characters as paragraph breaks.

    Args:
        doc: The Document instance.
        heading: Sub-heading string.
        text: The long-form text content (may contain newlines).
    """
    t = _active_theme
    doc.add_heading(heading, level=2)

    if text and text.strip():
        for paragraph_text in text.strip().split("\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(paragraph_text.strip())
                run.font.size = Pt(t.size_table)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No information provided.")
        run.font.size = Pt(t.size_table)
        run.italic = True
        run.font.color.rgb = t.color_muted

    doc.add_paragraph("")


def bullet_list(doc, heading, items_str):
    """
    Add a sub-heading followed by a bulleted list.
    Items are expected as a newline-separated string.

    Args:
        doc: The Document instance.
        heading: Sub-heading string.
        items_str: Newline-separated string of list items.
    """
    t = _active_theme
    doc.add_heading(heading, level=2)

    if items_str and items_str.strip():
        items = [item.strip() for item in items_str.split("\n") if item.strip()]
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(t.size_table)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No items listed.")
        run.font.size = Pt(t.size_table)
        run.italic = True
        run.font.color.rgb = t.color_muted

    doc.add_paragraph("")


def signatures(doc, labels):
    """
    Add a signature block with underlines and labels in a grid layout.
    Labels are arranged in pairs (two per row).

    Args:
        doc: The Document instance.
        labels: List of signature label strings (e.g. ["Employee Signature",
                "Date", "HR Representative", "Date"]).
    """
    t = _active_theme
    doc.add_paragraph("")
    doc.add_heading("Signatures", level=1)

    num_rows = (len(labels) + 1) // 2
    num_cols = min(len(labels), 2)
    sig_table = doc.add_table(rows=num_rows, cols=num_cols)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_margins(sig_table, top=0.02, bottom=0.02)

    for i, label in enumerate(labels):
        cell = sig_table.rows[i // 2].cells[i % 2]
        p = cell.paragraphs[0]
        p.add_run("\n\n")
        p.add_run("_" * 35 + "\n")
        run = p.add_run(label)
        run.font.name = t.font_caption
        run.font.size = Pt(t.size_caption)
        run.font.color.rgb = t.color_muted


def footer(doc):
    """
    Add the standard FormForge auto-generated footer paragraph.

    Args:
        doc: The Document instance.
    """
    t = _active_theme
    doc.add_paragraph("")
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "This document was auto-generated by FormForge. "
        "Please review all information for accuracy."
    )
    fr.font.name = t.font_caption
    fr.font.size = Pt(t.size_footer)
    fr.font.color.rgb = t.color_footer
    fr.italic = True


def address(doc, heading, raw_json):
    """
    Add a formatted mailing address block under a heading.
    Parses a JSON string with keys: street, city, state, zip.
    Handles empty/missing components gracefully.

    Args:
        doc: The Document instance.
        heading: Sub-heading string.
        raw_json: JSON string with address fields, or "{}".
    """
    t = _active_theme
    doc.add_heading(heading, level=2)

    try:
        addr = json.loads(raw_json) if raw_json else {}
    except (json.JSONDecodeError, TypeError):
        addr = {}

    street = addr.get("street", "").strip()
    if street:
        p = doc.add_paragraph()
        r = p.add_run(street + "\n")
        r.font.size = Pt(t.size_table)

        # Build city/state/zip line, skipping empty parts
        parts = []
        city = addr.get("city", "").strip()
        state = addr.get("state", "").strip()
        zipcode = addr.get("zip", "").strip()
        if city:
            parts.append(city)
        if state:
            city_state = ", ".join(parts + [state]) if parts else state
            parts = [city_state]
        if zipcode:
            parts.append(zipcode)
        city_line = " ".join(parts)
        if city_line:
            r = p.add_run(city_line)
            r.font.size = Pt(t.size_table)
    else:
        p = doc.add_paragraph()
        r = p.add_run("No address provided.")
        r.italic = True
        r.font.color.rgb = t.color_muted

    doc.add_paragraph("")


def image(doc, b64_str, width_inches=3.0, placeholder="No image uploaded."):
    """
    Embed a base64 data-URI image or render a placeholder if absent/invalid.

    Args:
        doc: The Document instance.
        b64_str: Base64 data URI string (e.g. "data:image/png;base64,...") or "".
        width_inches: Image width in inches (default: 3.0).
        placeholder: Text to show when no image is available.
    """
    if b64_str and "," in b64_str:
        try:
            img_data = b64_str.split(",")[1]
            img_bytes = base64.b64decode(img_data)
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
            return
        except (ValueError, binascii.Error, OSError, UnrecognizedImageError):
            pass

    p = doc.add_paragraph()
    r = p.add_run(placeholder)
    r.italic = True
    r.font.color.rgb = _active_theme.color_muted


def signature(doc, b64_str, label, width_inches=2.5):
    """
    Add a signature image (or placeholder underline) with a label beneath.

    Args:
        doc: The Document instance.
        b64_str: Base64 data URI string or "".
        label: Label text below the signature (e.g. "Employee Signature").
        width_inches: Signature image width in inches (default: 2.5).
    """
    t = _active_theme
    p = doc.add_paragraph()
    if b64_str and "," in b64_str:
        try:
            sig_data = b64_str.split(",")[1]
            sig_bytes = base64.b64decode(sig_data)
            doc.add_picture(io.BytesIO(sig_bytes), width=Inches(width_inches))
        except (ValueError, binascii.Error, OSError, UnrecognizedImageError):
            p.add_run("_" * 40)
    else:
        p.add_run("_" * 40)
    lp = doc.add_paragraph()
    lr = lp.add_run(label)
    lr.font.name = t.font_caption
    lr.font.size = Pt(t.size_caption)
    lr.font.color.rgb = t.color_muted


def repeater_table(doc, headers, items, field_keys, currency_keys=None):
    """
    Render a repeater field as a headed table with a shaded header row.

    Args:
        doc: The Document instance.
        headers: List of column header strings.
        items: List of dicts (parsed repeater JSON).
        field_keys: List of dict keys corresponding to each column.
        currency_keys: Optional set/list of field_keys that should be
                       formatted as currency (e.g. ["amount", "unit_cost"]).
    """
    if currency_keys is None:
        currency_keys = set()
    else:
        currency_keys = set(currency_keys)

    t = _active_theme

    if items:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        border_color = str(t.color_muted)
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(
            parse_xml(
                f"<w:tblBorders {nsdecls('w')}>"
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f"</w:tblBorders>"
            )
        )
        _set_cell_margins(table, top=0.02, bottom=0.02)

        header_row = table.rows[0]
        _shade_cells(header_row, str(t.color_accent))
        for i, header_text in enumerate(headers):
            cell_p = header_row.cells[i].paragraphs[0]
            r = cell_p.add_run(header_text)
            r.font.name = t.font_heading
            r.font.size = Pt(t.size_table)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for item in items:
            row = table.add_row()
            for col_idx, key in enumerate(field_keys):
                val = item.get(key, "")
                if key in currency_keys:
                    try:
                        val = f"${float(val):,.2f}"
                    except (ValueError, TypeError):
                        val = str(val)
                cell_p = row.cells[col_idx].paragraphs[0]
                r = cell_p.add_run(str(val))
                r.font.size = Pt(t.size_table)
    else:
        p = doc.add_paragraph()
        r = p.add_run("No line items provided.")
        r.italic = True
        r.font.color.rgb = t.color_muted

    doc.add_paragraph("")


def finalize(doc):
    """
    Serialize a Document to bytes.

    Args:
        doc: The Document instance.

    Returns:
        bytes: The .docx file as raw bytes.
    """
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
