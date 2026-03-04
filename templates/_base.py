"""
FormForge Shared Template Base
==============================
Shared helper functions for all FormForge DOCX templates.
Works in both standard Python (import _base) and Pyodide (via exec()).

Usage in templates:
    import _base
    doc = _base.new_doc("My Form Title", "Subtitle here")
    _base.add_table_section(doc, "Section Name", [("Label", "Value"), ...])
    _base.add_longtext(doc, "Notes", some_text)
    _base.add_bullet_list(doc, "Skills", newline_separated_str)
    _base.add_signatures(doc, ["Signer 1", "Signer 2"])
    return _base.finalize(doc)
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# -- Standard color palette --
COLOR_DARK_NAVY = RGBColor(0x1A, 0x1A, 0x3E)
COLOR_MEDIUM_BLUE = RGBColor(0x33, 0x33, 0x66)
COLOR_SOFT_BLUE = RGBColor(0x66, 0x66, 0x99)
COLOR_MUTED = RGBColor(0x99, 0x99, 0x99)
COLOR_LIGHT_MUTED = RGBColor(0xAA, 0xAA, 0xAA)


def new_doc(title_text, subtitle_text="", font_name="Calibri", font_size=11):
    """
    Create a styled Document with a centered title and optional subtitle.

    Args:
        title_text: Main heading text.
        subtitle_text: Optional subtitle displayed below the title.
        font_name: Base font for the document (default: "Calibri").
        font_size: Base font size in points (default: 11).

    Returns:
        A python-docx Document instance.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)

    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLOR_MEDIUM_BLUE

    doc.add_paragraph("")

    if subtitle_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle_text)
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_SOFT_BLUE

        doc.add_paragraph("")

    return doc


def add_table_section(doc, heading, rows):
    """
    Add a heading followed by a two-column key/value table.

    Args:
        doc: The Document instance.
        heading: Section heading string.
        rows: List of (label, value) tuples.
    """
    doc.add_heading(heading, level=1)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    for label, value in rows:
        row = table.add_row()

        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)

        vp = row.cells[1].paragraphs[0]
        display_value = str(value) if value else "\u2014"
        vr = vp.add_run(display_value)
        vr.font.size = Pt(10)

    doc.add_paragraph("")


def add_longtext(doc, heading, text):
    """
    Add a sub-heading followed by one or more paragraphs of body text.
    Respects newline characters as paragraph breaks.

    Args:
        doc: The Document instance.
        heading: Sub-heading string.
        text: The long-form text content (may contain newlines).
    """
    doc.add_heading(heading, level=2)

    if text and text.strip():
        for paragraph_text in text.strip().split("\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(paragraph_text.strip())
                run.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No information provided.")
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph("")


def add_bullet_list(doc, heading, items_str):
    """
    Add a sub-heading followed by a bulleted list.
    Items are expected as a newline-separated string.

    Args:
        doc: The Document instance.
        heading: Sub-heading string.
        items_str: Newline-separated string of list items.
    """
    doc.add_heading(heading, level=2)

    if items_str and items_str.strip():
        items = [item.strip() for item in items_str.split("\n") if item.strip()]
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No items listed.")
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph("")


def add_signatures(doc, labels):
    """
    Add a signature block with underlines and labels in a grid layout.
    Labels are arranged in pairs (two per row).

    Args:
        doc: The Document instance.
        labels: List of signature label strings (e.g. ["Employee Signature",
                "Date", "HR Representative", "Date"]).
    """
    doc.add_paragraph("")
    doc.add_heading("Signatures", level=1)

    num_rows = (len(labels) + 1) // 2
    num_cols = min(len(labels), 2)
    sig_table = doc.add_table(rows=num_rows, cols=num_cols)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, label in enumerate(labels):
        cell = sig_table.rows[i // 2].cells[i % 2]
        p = cell.paragraphs[0]
        p.add_run("\n\n")
        p.add_run("_" * 35 + "\n")
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED


def finalize(doc):
    """
    Serialize a Document to bytes.

    Args:
        doc: The Document instance.

    Returns:
        bytes: The .docx file as raw bytes.
    """
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
