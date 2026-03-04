"""
FormForge DOCX Template: Expense Report
========================================
Demonstrates all new field types (number, currency, heading, hidden,
address, file, signature, repeater) alongside existing types.

Field type notes:
  - text/email/tel/date/select/radio/textarea → str
  - number/currency → str (e.g. "42.5", "1250.00")
  - hidden → str (static default_value from schema)
  - heading → not passed (skipped in collectFormData)
  - checkbox → comma-separated str
  - longtext → str with possible newlines
  - list → newline-separated str
  - address → JSON string: {"street","city","state","zip"}
  - file/signature → base64 data URI str or ""
  - repeater → JSON array string: [{"field":"value"}, ...]
"""

import io
import json
import base64
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import _base


def generate_docx(data):
    """
    Generate an Expense Report Document from form data.

    Args:
        data: dict mapping field IDs to their string values.

    Returns:
        bytes: The generated .docx file as raw bytes.
    """
    name = data.get("employee_name", "")
    dept = data.get("department", "")
    report_date = data.get("report_date", "")
    form_ver = data.get("form_version", "")

    doc = _base.new_doc(
        "Expense Report",
        f"{name} — {dept} — {report_date}",
    )

    # ── Report Details ──────────────────────────────────────
    total = data.get("total_amount", "0")
    try:
        total_fmt = f"${float(total):,.2f}"
    except (ValueError, TypeError):
        total_fmt = f"${total}"

    _base.add_table_section(
        doc,
        "Report Details",
        [
            ("Employee", name),
            ("Department", dept),
            ("Report Date", report_date),
            ("Total Amount", total_fmt),
            ("Form Version", form_ver),
        ],
    )

    # ── Line Items (repeater) ───────────────────────────────
    doc.add_heading("Expense Line Items", level=1)

    raw_items = data.get("line_items", "[]")
    try:
        items = json.loads(raw_items)
    except (json.JSONDecodeError, TypeError):
        items = []

    if items:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"
        for i, header in enumerate(["Description", "Amount", "Category"]):
            cell_p = table.rows[0].cells[i].paragraphs[0]
            r = cell_p.add_run(header)
            r.bold = True
            r.font.size = Pt(10)
        for item in items:
            row = table.add_row()
            row.cells[0].text = item.get("description", "")
            amt = item.get("amount", "0")
            try:
                row.cells[1].text = f"${float(amt):,.2f}"
            except (ValueError, TypeError):
                row.cells[1].text = str(amt)
            row.cells[2].text = item.get("category", "")
    else:
        p = doc.add_paragraph()
        r = p.add_run("No line items provided.")
        r.italic = True
        r.font.color.rgb = _base.COLOR_MUTED

    doc.add_paragraph("")

    # ── Number of Receipts (number field) ───────────────────
    receipt_count = data.get("item_count", "0")
    p = doc.add_paragraph()
    r = p.add_run(f"Number of receipts attached: {receipt_count}")
    r.font.size = Pt(10)

    doc.add_paragraph("")

    # ── Receipt Photo (file field) ──────────────────────────
    doc.add_heading("Receipt / Invoice", level=1)
    receipt_b64 = data.get("receipt_photo", "")
    if receipt_b64 and "," in receipt_b64:
        try:
            img_data = receipt_b64.split(",")[1]
            img_bytes = base64.b64decode(img_data)
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(3))
        except Exception:
            p = doc.add_paragraph()
            r = p.add_run("[Receipt image could not be embedded]")
            r.italic = True
            r.font.color.rgb = _base.COLOR_MUTED
    else:
        p = doc.add_paragraph()
        r = p.add_run("No receipt uploaded.")
        r.italic = True
        r.font.color.rgb = _base.COLOR_MUTED

    doc.add_paragraph("")

    # ── Notes ───────────────────────────────────────────────
    notes = data.get("notes", "")
    if notes and notes.strip():
        doc.add_heading("Additional Notes", level=1)
        p = doc.add_paragraph()
        r = p.add_run(notes)
        r.font.size = Pt(10)
        doc.add_paragraph("")

    # ── Mailing Address (address field) ─────────────────────
    doc.add_heading("Reimbursement Mailing Address", level=1)
    raw_addr = data.get("mailing_address", "{}")
    try:
        addr = json.loads(raw_addr)
    except (json.JSONDecodeError, TypeError):
        addr = {}

    if addr.get("street"):
        p = doc.add_paragraph()
        lines = [
            addr.get("street", ""),
            f"{addr.get('city', '')}, {addr.get('state', '')} {addr.get('zip', '')}".strip(),
        ]
        for line in lines:
            if line.strip() and line.strip() != ",":
                r = p.add_run(line.strip() + "\n")
                r.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        r = p.add_run("No address provided.")
        r.italic = True
        r.font.color.rgb = _base.COLOR_MUTED

    doc.add_paragraph("")

    # ── Signatures (signature fields) ───────────────────────
    doc.add_heading("Approval Signatures", level=1)

    for label, field_id in [
        ("Employee Signature", "employee_signature"),
        ("Manager Signature", "manager_signature"),
    ]:
        sig_b64 = data.get(field_id, "")
        p = doc.add_paragraph()
        if sig_b64 and "," in sig_b64:
            try:
                sig_data = sig_b64.split(",")[1]
                sig_bytes = base64.b64decode(sig_data)
                doc.add_picture(io.BytesIO(sig_bytes), width=Inches(2.5))
            except Exception:
                r = p.add_run("_" * 40)
        else:
            r = p.add_run("_" * 40)
        lp = doc.add_paragraph()
        lr = lp.add_run(label)
        lr.font.size = Pt(9)
        lr.font.color.rgb = _base.COLOR_MUTED

    doc.add_paragraph("")

    # ── Footer ──────────────────────────────────────────────
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "This document was auto-generated by FormForge. "
        "Please review all information for accuracy."
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = _base.COLOR_MUTED
    fr.italic = True

    return _base.finalize(doc)
