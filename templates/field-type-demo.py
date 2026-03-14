"""
FormForge DOCX Template: Event Registration (Field Type Demo)
=============================================================
Demonstrates all 24 field types, wizard mode, and conditional visibility.

Field type value formats:
  - text/email/tel/date/select/radio/textarea → str
  - number/currency → str (e.g. "50", "1250.00")
  - hidden → str (static default_value from schema)
  - heading → not passed (skipped in collectFormData)
  - info → not passed (skipped in collectFormData)
  - checkbox → comma-separated str
  - longtext → str with possible newlines
  - list → newline-separated str
  - address → JSON string: {"street","city","state","zip"}
  - file/signature → base64 data URI str or ""
  - repeater → JSON array string: [{"field":"value"}, ...]
  - time → str (e.g. "09:00")
  - url → str (e.g. "https://example.com")
  - toggle → str ("true" or "false")
  - datetime → str (e.g. "2026-06-15T14:30")
  - multi_select → comma-separated str (e.g. "AI/ML, Cloud")
"""

import json
from docx.shared import Pt

import stencils


def generate_docx(data: dict[str, str]) -> bytes:
    """
    Generate an Event Registration document from form data.

    Args:
        data: dict mapping field IDs to their string values.

    Returns:
        bytes: The generated .docx file as raw bytes.
    """
    stencils.set_theme(stencils.THEME_MINIMAL)

    name = data.get("full_name", "")
    email = data.get("email", "")
    event_date = data.get("event_date", "")
    form_ver = data.get("form_version", "")

    doc = stencils.new_doc(
        "Event Registration",
        f"{name} — {event_date}",
    )

    # ── Step 1: Applicant Info ────────────────────────────
    stencils.table_section(
        doc,
        "Applicant Info",
        [
            ("Full Name", name),
            ("Email", email),
            ("Phone", data.get("phone", "")),
            ("Preferred Date", event_date),
            ("Preferred Time", stencils.format_time(data.get("event_time", ""))),
            ("Website", data.get("website", "")),
            ("Form Version", form_ver),
        ],
    )

    # ── Step 2: Event Details ─────────────────────────────
    event_type = data.get("event_type", "")
    attendance = data.get("attendance_mode", "")

    stencils.table_section(
        doc,
        "Event Details",
        [
            ("Event Type", event_type),
            ("Attendance Mode", attendance),
        ],
    )

    # textarea — brief description
    description = data.get("event_description", "")
    if description and description.strip():
        doc.add_heading("Brief Description", level=2)
        p = doc.add_paragraph()
        r = p.add_run(description)
        r.font.size = Pt(10)
        doc.add_paragraph("")

    # longtext — detailed proposal
    proposal = data.get("detailed_proposal", "")
    if proposal and proposal.strip():
        stencils.longtext(doc, "Detailed Proposal", proposal)

    # address (conditional — in-person only)
    raw_addr = data.get("venue_address", "{}")
    try:
        addr = json.loads(raw_addr) if raw_addr else {}
    except (json.JSONDecodeError, TypeError):
        addr = {}

    if addr.get("street"):
        stencils.address(doc, "Venue Address", raw_addr)

    # checkbox (conditional — in-person dietary restrictions)
    dietary = data.get("dietary_restrictions", "")
    if dietary and dietary.strip():
        dietary_nl = "\n".join(d.strip() for d in dietary.split(",") if d.strip())
        stencils.bullet_list(doc, "Dietary Restrictions", dietary_nl)

    # text (conditional — virtual platform)
    virtual_platform = data.get("virtual_platform", "")
    if virtual_platform and virtual_platform.strip():
        doc.add_heading("Virtual Platform", level=2)
        p = doc.add_paragraph()
        r = p.add_run(virtual_platform)
        r.font.size = Pt(10)
        doc.add_paragraph("")

    # multi_select — session topics
    topics = data.get("session_topics", "")
    if topics and topics.strip():
        topics_nl = "\n".join(t.strip() for t in topics.split(",") if t.strip())
        stencils.bullet_list(doc, "Session Topics", topics_nl)

    # ── Step 3: Budget & Items ────────────────────────────
    budget = data.get("budget_amount", "0")
    try:
        budget_fmt = f"${float(budget):,.2f}"
    except (ValueError, TypeError):
        budget_fmt = f"${budget}"

    attendees = data.get("attendee_count", "")

    stencils.table_section(
        doc,
        "Budget & Logistics",
        [
            ("Estimated Budget", budget_fmt),
            ("Expected Attendees", attendees),
        ],
    )

    # repeater — expense line items
    doc.add_heading("Budget Line Items", level=1)
    raw_items = data.get("expense_items", "[]")
    try:
        items = json.loads(raw_items)
    except (json.JSONDecodeError, TypeError):
        items = []

    stencils.repeater_table(
        doc,
        headers=["Item", "Qty", "Unit Cost", "Category"],
        items=items,
        field_keys=["item_name", "quantity", "unit_cost", "category"],
        currency_keys=["unit_cost"],
    )

    # list — special requests
    special = data.get("special_requests", "")
    if special and special.strip():
        stencils.bullet_list(doc, "Special Requests", special)

    # datetime — submission deadline
    deadline = data.get("submission_deadline", "")
    if deadline and deadline.strip():
        # Format datetime-local value (e.g. "2026-05-01T17:00" → "2026-05-01 at 5:00 PM")
        deadline_display = stencils.format_time(deadline)
        stencils.table_section(
            doc,
            "Deadline",
            [("Proposal Submission Deadline", deadline_display)],
        )

    # ── Step 4: Attachments & Approval ────────────────────

    # file — supporting document
    doc.add_heading("Supporting Document", level=1)
    stencils.image(
        doc,
        data.get("supporting_doc", ""),
        width_inches=3.0,
        placeholder="No document uploaded.",
    )
    doc.add_paragraph("")

    # textarea — additional notes
    notes = data.get("additional_notes", "")
    if notes and notes.strip():
        stencils.longtext(doc, "Additional Notes", notes)

    # toggle — terms agreement
    agree = data.get("agree_terms", "false")
    stencils.table_section(
        doc,
        "Terms",
        [("Agreed to Terms & Conditions", "Yes" if agree == "true" else "No")],
    )

    # signature — applicant signature
    doc.add_heading("Applicant Signature", level=1)
    stencils.signature(doc, data.get("applicant_signature", ""), "Applicant Signature")
    doc.add_paragraph("")

    # ── Footer ────────────────────────────────────────────
    stencils.footer(doc)

    return stencils.finalize(doc)
