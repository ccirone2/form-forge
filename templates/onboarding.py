"""
FormForge DOCX Template: Employee Onboarding Document
======================================================
This script is loaded by FormForge's Pyodide runtime.
It must export a `generate_docx(data)` function that accepts
a dict of form field values and returns docx file bytes.

The `data` dict keys match the field `id` values in the
corresponding schema JSON (schemas/onboarding.json).

Field type notes:
  - text/email/tel/date/select/radio/textarea → str
  - checkbox → comma-separated str (e.g. "GitHub, Jira, Figma")
  - longtext → str with possible newlines for paragraphs
  - list → newline-separated str (e.g. "Python\\nJavaScript\\nRust")
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# -- Color palette --
COLOR_DARK_NAVY = RGBColor(0x1A, 0x1A, 0x3E)
COLOR_MEDIUM_BLUE = RGBColor(0x33, 0x33, 0x66)
COLOR_SOFT_BLUE = RGBColor(0x66, 0x66, 0x99)
COLOR_MUTED = RGBColor(0x99, 0x99, 0x99)
COLOR_LIGHT_MUTED = RGBColor(0xAA, 0xAA, 0xAA)


def generate_docx(data):
    """
    Generate an Employee Onboarding Document from form data.

    Args:
        data: dict mapping field IDs to their string values.

    Returns:
        bytes: The generated .docx file as raw bytes.
    """
    doc = Document()

    # ── Global styles ──────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──────────────────────────────────────────────────
    title = doc.add_heading("Employee Onboarding Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLOR_MEDIUM_BLUE

    doc.add_paragraph("")

    # ── Summary line ───────────────────────────────────────────
    first = data.get("first_name", "")
    last = data.get("last_name", "")
    start = data.get("start_date", "TBD")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Prepared for {first} {last} — Start Date: {start}")
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_SOFT_BLUE

    doc.add_paragraph("")

    # ── Section: Personal Information ──────────────────────────
    _add_table_section(doc, "Personal Information", [
        ("First Name",     data.get("first_name", "")),
        ("Last Name",      data.get("last_name", "")),
        ("Email",          data.get("email", "")),
        ("Phone",          data.get("phone", "")),
        ("Date of Birth",  data.get("date_of_birth", "")),
        ("Start Date",     data.get("start_date", "")),
    ])

    # ── Section: Role & Department ─────────────────────────────
    _add_table_section(doc, "Role & Department", [
        ("Department",        data.get("department", "")),
        ("Job Title",         data.get("job_title", "")),
        ("Employment Type",   data.get("employment_type", "")),
        ("Reporting Manager", data.get("manager", "")),
    ])

    # ── Section: Equipment & Access ────────────────────────────
    equipment = data.get("equipment_needs", "")
    software = data.get("software_access", "")

    _add_table_section(doc, "Equipment & Access", [
        ("Laptop",               data.get("laptop_preference", "")),
        ("Additional Equipment", equipment if equipment else "None requested"),
        ("Software Access",      software if software else "None requested"),
    ])

    # ── Section: Skills & Experience ───────────────────────────
    doc.add_heading("Skills & Experience", level=1)

    _add_longtext(doc, "Professional Bio", data.get("bio", ""))
    _add_bullet_list(doc, "Key Skills", data.get("skills", ""))
    _add_bullet_list(doc, "Certifications & Licenses", data.get("certifications", ""))
    _add_longtext(doc, "Notable Prior Projects", data.get("prior_projects", ""))

    # ── Section: Additional Information ────────────────────────
    _add_table_section(doc, "Additional Information", [
        ("Dietary Restrictions", data.get("dietary_restrictions", "")),
        ("Emergency Contact",    data.get("emergency_contact", "")),
        ("Emergency Phone",      data.get("emergency_phone", "")),
        ("Notes",                data.get("notes", "")),
    ])

    _add_bullet_list(doc, "First 90 Days Goals", data.get("onboarding_goals", ""))

    # ── Signatures ─────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_heading("Signatures", level=1)

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_labels = [
        "Employee Signature", "Date",
        "HR Representative",  "Date",
    ]
    for i, label in enumerate(sig_labels):
        cell = sig_table.rows[i // 2].cells[i % 2]
        p = cell.paragraphs[0]
        p.add_run("\n\n")
        p.add_run("_" * 35 + "\n")
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED

    # ── Footer ─────────────────────────────────────────────────
    doc.add_paragraph("")
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "This document was auto-generated by FormForge. "
        "Please review all information for accuracy."
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = COLOR_LIGHT_MUTED
    fr.italic = True

    # ── Serialize ──────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def _add_table_section(doc, title_text, fields):
    """
    Add a heading followed by a two-column label/value table.

    Args:
        doc: The Document instance.
        title_text: Section heading string.
        fields: List of (label, value) tuples.
    """
    doc.add_heading(title_text, level=1)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    for label, value in fields:
        row = table.add_row()

        # Label cell
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)

        # Value cell
        vp = row.cells[1].paragraphs[0]
        display_value = str(value) if value else "\u2014"
        vr = vp.add_run(display_value)
        vr.font.size = Pt(10)

    doc.add_paragraph("")


def _add_longtext(doc, title_text, text):
    """
    Add a sub-heading followed by one or more paragraphs of body text.
    Respects newline characters as paragraph breaks.

    Args:
        doc: The Document instance.
        title_text: Sub-heading string.
        text: The long-form text content (may contain newlines).
    """
    doc.add_heading(title_text, level=2)

    if text and text.strip():
        for paragraph_text in text.strip().split("\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(paragraph_text.strip())
                run.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No information provided.")
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph("")


def _add_bullet_list(doc, title_text, items_str):
    """
    Add a sub-heading followed by a bulleted list.
    Items are expected as a newline-separated string.

    Args:
        doc: The Document instance.
        title_text: Sub-heading string.
        items_str: Newline-separated string of list items.
    """
    doc.add_heading(title_text, level=2)

    if items_str and items_str.strip():
        items = [item.strip() for item in items_str.split("\n") if item.strip()]
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No items listed.")
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph("")
